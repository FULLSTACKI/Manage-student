from src.domain.repositories.export_file import IsExportFile
import pandas as pd
import os
import io
from docx import Document
from docx.shared import Inches

class DocxExportFile(IsExportFile):
    def export(self, reports):
        try: 
            df = pd.DataFrame(reports.get("data_table"))
            output_buffer = io.BytesIO()
            doc = Document()
            doc.add_heading("Báo cáo dữ liệu:")
            
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
                    
            chart_path = reports.get("chart_path")
            if chart_path and os.path.exists(chart_path):
                doc.add_heading("Trực quan hóa dữ liệu:")
                doc.add_picture(chart_path, width=Inches(6))
            
            insight_text = reports.get("response")
            if insight_text:
                doc.add_paragraph(insight_text)
            
            doc.save(output_buffer)
            output_buffer.seek(0)
            return output_buffer
        except Exception as e:
            raise e